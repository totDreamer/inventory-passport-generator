import warnings
import re

# Сохраняем оригинальную функцию
original_showwarning = warnings.showwarning

# Свои условия
def custom_showwarning(message, category, filename, lineno, file=None, line=None):
    msg = str(message)
    if (
        re.search(r"pkg_resources is deprecated as an API", msg)
        or re.search(r"Data Validation extension is not supported", msg)
    ):
        return  # Не выводим эти два предупреждения
    return original_showwarning(message, category, filename, lineno, file, line)

warnings.showwarning = custom_showwarning


import os
import pandas as pd
import FreeSimpleGUI as sg
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from copy import deepcopy

def set_vertical_text(cell):
    """Устанавливает вертикальный текст снизу вверх."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), 'btLr')  # btLr — снизу вверх
    tcPr.append(textDirection)

def generate_inventory_report(excel_path, template_path, output_path, start_row, end_row):
    try:
        # Загрузка данных из Excel
        df = pd.read_excel(excel_path, skiprows=5, dtype=str)
        df.columns = df.columns.str.strip()
        df = df.fillna("")

        # Открываем шаблон Word
        doc = Document(template_path)
        table = doc.tables[0]

        # Соответствие колонок Excel и ячеек в Word
        column_mapping = {
            'organisation': 2, 'cabinet': 3, 'user_name': 1, 'department': 4,
            'pc_mark': 5, 'pc_model': 5, 'pc_serial_number': 6, 'pc_inv_number': 7,
            'pc_name': 8, 'login': 9, 'domain': 10, 'ip_addr': 11,
            'os': 12, 'cpu_model': 13, 'cpu_freq': 13, 'ddr_size': 14, 'hdd_mark': 15,'hdd_size': 15, 'hdd2_mark': 15,'hdd2_size': 15,
            'ssd_mark': 15,'ssd_size': 15, 'monitor_mark': 16, 'monitor_model': 16, 'monitor_inv_num': 17, 'monitor_sn': 18,
            'printer_dev': 19, 'printer_model': 19, 'printer_color': 20, 'printer_sn': 21, 'printer_inv_num': 22,
            'ip_dev': 23, 'ip_model': 23, 'ip_sn': 24, 'ip_inv': 26, 'antivirus': 27, 'dlp': 28
        }

        # Заполнение таблицы (только выбранные строки)
        for idx, row in df.iloc[start_row-1:end_row].iterrows():
            new_row = table.add_row()
            new_row.cells[0].text = str(len(table.rows) - 3)
            
            for excel_col, word_cell_idx in column_mapping.items():
                if excel_col in row and row[excel_col]:                     # Проверка на наличие ключа и имеющегося значения по нему
                    if word_cell_idx in [5, 16, 19, 23]:                    # Объединение данных для ячеек, в которые нужно вписать несколько значений
                        existing_text = new_row.cells[word_cell_idx].text
                        if existing_text and row[excel_col] != '-':
                            new_text = f"{existing_text} {row[excel_col]}"  # Через пробел
                        else:
                            new_text = row[excel_col]
                        if word_cell_idx == 23:                             # Обработка добавления прочерка для mac-адреса, если заполняются данные по телефону
                            new_row.cells[25].text = '-'
                        new_row.cells[word_cell_idx].text = new_text.strip()
                    elif word_cell_idx == 13:                               # Объединение для cpu_model и cpu_freq по шаблону
                        existing_text = new_row.cells[13].text
                        if existing_text:
                            new_text = f"{existing_text}, {row[excel_col]}GHz"  # Через запятую и с добавлением "GHz"
                        else:
                            new_text = row[excel_col]
                        new_row.cells[13].text = new_text.strip()
                    elif word_cell_idx == 15:                               # Объединение для hdd и ssd ячейки по шаблону
                        existing_text = new_row.cells[15].text
                        if existing_text:
                            if excel_col in ('hdd_size, hdd2_size'):
                                new_text = f"{existing_text} {row[excel_col]}GB" 
                            elif excel_col == 'ssd_mark':  
                                new_text = f"{existing_text}, ({row[excel_col]}"
                            elif excel_col == 'ssd_size':
                                new_text = f"{existing_text} {row[excel_col]}GB)"
                        else:
                            if excel_col == 'ssd_mark':                     # Обработка случая, где есть SSD, но нет HDD, чтоб не терялась левая скобка
                                new_text = f'({row[excel_col]}'
                            else:
                                new_text = row[excel_col]
                        new_row.cells[15].text = new_text.strip()
                    else:                                                       # Добавление данных для остальных столбцов
                        new_row.cells[word_cell_idx].text = str(row[excel_col])

        # Форматирование таблицы (начиная с 3-й строки)
        for row_idx, row in enumerate(table.rows):
            if row_idx >= 2:  # Индексация с 0, поэтому 2 - это третья строка
                # Установка фиксированной высоты строки
                row.height = Inches(1.4)  # Фиксированная высота (0.4 дюйма ≈ 1 см)
                
                # Форматирование текста в ячейках
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(8)
                    if cell != row.cells[0]:
                        set_vertical_text(cell)

        # Сохранение файла
        doc.save(output_path)
        return True, f"Инвентаризационный отчёт успешно сохранён: {output_path}"
    
    except Exception as e:
        return False, f"Ошибка при создании инвентаризационного отчёта: {str(e)}"

def generate_passports(excel_path, template_path, output_path, start_row, end_row):
    try:
        df = pd.read_excel(excel_path, skiprows=5, dtype=str)
        df.columns = df.columns.str.strip()
        df = df.fillna("")
        rows = df.iloc[start_row - 1:end_row].to_dict(orient="records")

        base_template = DocxTemplate(template_path)
        base_template.render({})
        base_template.save("~temp_base.docx")

        final_doc = Document("~temp_base.docx")

        # Очищаем тело финального документа
        for element in final_doc.element.body[:]:
            final_doc.element.body.remove(element)

        generated_count = 0

        for i, context in enumerate(rows, start=1):
            try:
                template = DocxTemplate(template_path)
                template.render(context)
                template.save("~temp_passport.docx")

                temp_doc = Document("~temp_passport.docx")

                for elem in temp_doc.element.body:
                    # Исключаем пустые абзацы и служебные элементы (например, w:sectPr)
                    if elem.tag.endswith("sectPr"):
                        continue
                    if elem.tag.endswith("p"):  # Параграф
                        texts = [node.text for node in elem.iter() if node.text and node.text.strip()]
                        if not texts:
                            continue  # пропустить пустые параграфы

                    final_doc.element.body.append(deepcopy(elem))

                if i < len(rows):
                    final_doc.add_page_break()

                generated_count += 1

            except Exception as e:
                print(f"Ошибка при создании паспорта {i}: {str(e)}")

        final_doc.save(output_path)
        os.remove("~temp_base.docx")
        os.remove("~temp_passport.docx")

        return True, f"Сгенерировано {generated_count} паспортов в одном файле: {output_path}"

        # Сохраняем финальный файл
        final_doc.save(output_path)
        os.remove("~temp_base.docx")
        os.remove("~temp_passport.docx")
        return True, f"Сгенерировано {generated_count} паспортов в одном файле: {output_path}"

    except Exception as e:
        return False, f"Ошибка при создании паспортов: {str(e)}"

# GUI-интерфейс
sg.theme("LightGrey1")

layout = [
    [sg.Text("Excel-файл:"), sg.Input(key="-EXCEL-"), sg.FileBrowse(file_types=(("Excel Files", "*.xlsx"),))],
    
    [sg.Frame("Инвентаризационный отчёт", [
        [sg.Text("Шаблон Word:"), sg.Input(key="-INV_TEMPLATE-"), sg.FileBrowse(file_types=(("Word Files", "*.docx"),))],
        [sg.Text("Куда сохранить:"), sg.Input(key="-INV_OUTPUT-"), sg.FileSaveAs(file_types=(("Word Files", "*.docx"),))],
    ])],
    
    [sg.Frame("Паспорта оборудования", [
        [sg.Text("Шаблон Word:"), sg.Input(key="-PASSP_TEMPLATE-"), sg.FileBrowse(file_types=(("Word Files", "*.docx"),))],
        [sg.Text("Куда сохранить:"), sg.Input(key="-PASSP_OUTPUT-"), sg.FileSaveAs(file_types=(("Word Files", "*.docx"),))],
    ])],
    
    [sg.Text("Строки (от/до):"), sg.Input("1", size=5, key="-START-"), sg.Input("5", size=5, key="-END-")],
    [sg.Checkbox("Генерировать инвентаризационный отчёт", default=True, key="-GEN_INV-"), 
     sg.Checkbox("Генерировать паспорта", default=True, key="-GEN_PASSP-")],
    [sg.Button("Сгенерировать"), sg.Button("Выход")],
    [sg.Output(size=(80, 10), key="-LOG-")]
]

window = sg.Window("Генератор отчётов и паспортов", layout)

while True:
    event, values = window.read()
    if event in (None, "Выход"):
        break
    if event == "Сгенерировать":
        # Проверка ввода
        if not values["-EXCEL-"]:
            sg.popup_error("Укажите Excel-файл!")
            continue
        
        try:
            start_row = int(values["-START-"])
            end_row = int(values["-END-"])
        except ValueError:
            sg.popup_error("Номера строк должны быть числами!")
            continue
        
        # Проверка выбранных опций
        gen_inv = values["-GEN_INV-"]
        gen_passp = values["-GEN_PASSP-"]
        
        if not gen_inv and not gen_passp:
            sg.popup_error("Выберите хотя бы один тип отчёта для генерации!")
            continue
        
        # Проверка путей для выбранных опций
        if gen_inv and (not values["-INV_TEMPLATE-"] or not values["-INV_OUTPUT-"]):
            sg.popup_error("Для генерации инвентаризационного отчёта укажите шаблон и путь сохранения!")
            continue
        
        if gen_passp and (not values["-PASSP_TEMPLATE-"] or not values["-PASSP_OUTPUT-"]):
            sg.popup_error("Для генерации паспортов укажите шаблон и папку для сохранения!")
            continue
        
        # Запуск генерации
        print("="*50)
        print("Начало генерации отчётов...")
        
        if gen_inv:
            print("\nГенерация инвентаризационного отчёта:")
            success, message = generate_inventory_report(
                excel_path=values["-EXCEL-"],
                template_path=values["-INV_TEMPLATE-"],
                output_path=values["-INV_OUTPUT-"],
                start_row=start_row,
                end_row=end_row
            )
            print(message)
            if not success:
                sg.popup_error(message)
        
        if gen_passp:
            print("\nГенерация паспортов оборудования:")
            output_path = values["-PASSP_OUTPUT-"]
            success, message = generate_passports(
                excel_path=values["-EXCEL-"],
                template_path=values["-PASSP_TEMPLATE-"],
                output_path=output_path,
                start_row=start_row,
                end_row=end_row
            )

            print(message)
            if not success:
                sg.popup_error(message)
        
        print("\nГенерация завершена!")
        print("="*50)
        sg.popup_ok("Генерация отчётов завершена!")

window.close()